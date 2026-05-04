import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- 1. KONFIGURASI ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)

# --- 2. LOGIK AI (KRITERIA BERSEPADU) ---
def generate_advanced_plan_bm(topic, syllabus, extra_context):
    prompt = f"""
    Topik: {topic}. Kod Sukatan Pelajaran: {syllabus}. Konteks Tambahan: {extra_context}.
    Hasilkan rancangan pengajaran (lesson plan) profesional dalam BAHASA MELAYU sepenuhnya.
    
    Gunakan penanda (markers) yang TEPAT berikut untuk struktur dokumen:
    
    SECTION: TOPIK
    [Sila paparkan topik di sini mengikut input daripada pengguna]
    
    SECTION: OBJEKTIF PEMBELAJARAN
    [4 mata/point]
    
    SECTION: HASIL PEMBELAJARAN
    [4 mata/point]
    
    SECTION: KRITERIA KEJAYAAN
    [4 mata/point]
    
    SECTION: PENGETAHUAN SEDIA ADA
    [1 mata/point]
    
    SECTION: KATA KUNCI
    [6 item]
    
    SECTION: SOALAN TAHAP TINGGI (HOTS)
    [4 domain utama daripada Taksonomi Bloom]
    
    SECTION: KEWARGANEGARAAN DIGITAL
    [4 mata mengenai penggunaan teknologi beretika/Chromebooks/Canva/YouTube]

    SECTION: AKTIVITI PEMULA (INDUKSI)
    [Aktiviti 'Hook' dan pelan transisi]

    SECTION: STRATEGI PERBEZAAN (HIJAU)
    - HA (Pencapaian Tinggi): [1 aktiviti mencabar]

    SECTION: STRATEGI PERBEZAAN (KUNING)
    - MA (Pencapaian Sederhana): [1 aktiviti teras]

    SECTION: STRATEGI PERBEZAAN (MERAH)
    - LA (Pencapaian Rendah): [1 aktiviti berskala/scaffolded]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
    - Aktiviti 1: [Penerangan]
    - ----------------------------------------------------------------------------
    - Persediaan Guru: [Langkah demi langkah sebelum sesi]
    - ----------------------------------------------------------------------------
    - Objektif: [3 mata/point]
    - ----------------------------------------------------------------------------
    - Tugasan Pelajar: [Butiran langkah demi langkah]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
    - Aktiviti 2: [Penerangan]
    - -----------------------------------------------------------------------------
    - Persediaan Guru: [Langkah demi langkah sebelum sesi]
    - -----------------------------------------------------------------------------
    - Objektif: [3 mata/point]
    - -----------------------------------------------------------------------------
    - Tugasan Pelajar: [Butiran langkah demi langkah]
    
    SECTION: PLENARI (TIKET PENUTUP)
    [Aktiviti penutup 2-3 minit]

    SECTION: KERJA RUMAH
    [Tugasan berdasarkan topik]

    SECTION: CADANGAN LANGKAH SETERUSNYA
    - Aktiviti 'Hook' dan pelan transisi untuk sesi pengajaran hari berikutnya.
    
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Ralat Sistem: {str(e)}"

# --- 3. LOGIK EKSPORT WORD ---
def create_word_export_bm(topic, syllabus, text):
    doc = Document()
    doc.add_heading(f'Rancangan Mengajar : {topic}', 0)

    # Admin Header (Jadual Pentadbiran)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Minggu Ke:", "Tarikh:"], ["Jumlah Pelajar:", "Hari:"], ["Kelas:", "Tempoh:"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    # Parsing dan Kotak untuk SEMUA Seksyen
    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip()
        
        # Membersihkan teks daripada simbol markdown
        content = "\n".join(lines[1:]).strip().replace("**", "") 
        
        doc.add_heading(title.title(), level=1)
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = content
        doc.add_paragraph()
     
    # Kelulusan Ketua Jabatan (HOD)
    doc.add_page_break()
    doc.add_heading("Kelulusan & Ulasan Ketua Jabatan (HOD)", level=1)
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Ulasan:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Tarikh:"; hod_table.cell(1, 1).text = "Tandatangan:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. ANTARAMUKA PENGGUNA (GUI) ---
st.set_page_config(page_title="Sistem Rancangan Mengajar Versi 3.5", layout="wide")

st.title("🎓 LAMPIRAN RANCANGAN MENGAJAR PTE SENGKURONG")
st.info("Sila masukkan topik pelajaran, kod sukatan pelajaran subjek dan maklumat tambahan seperti Canva, YouTube atau infografik.")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Topik Pelajaran:")
with c2: u_syllabus = st.text_input("Kod Sukatan Pelajaran:")
u_extra = st.text_area("Konteks Tambahan (Pilihan):")

if st.button("🚀 JANA LAMPIRAN LENGKAP"):
    if u_topic and u_syllabus:
        with st.spinner("AI sedang menyelaraskan semua kriteria ke dalam rancangan anda..."):
            result = generate_advanced_plan_bm(u_topic, u_syllabus, u_extra)
            st.session_state['adv_plan_out_bm'] = result
    else:
        st.warning("Sila isi bahagian Topik dan Sukatan Pelajaran.")

if 'adv_plan_out_bm' in st.session_state:
    st.divider()
    st.subheader("Pratonton Draf AI")
    st.text_area("Kandungan", st.session_state['adv_plan_out_bm'], height=400)
    doc_file = create_word_export_bm(u_topic, u_syllabus, st.session_state['adv_plan_out_bm'])
    st.download_button("📥 Muat Turun Versi Word (.docx)", doc_file, f"PlanMengajar_{u_topic}.docx")

st.markdown("---")
st.caption("Sistem Rancangan Mengajar 3.0 | Penyedia: Hjh Nurul Haziqah Hj Nordin | © 2026 PTES Innovation")
